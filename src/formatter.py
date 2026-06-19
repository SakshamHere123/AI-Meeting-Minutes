import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.summarizer import MeetingMinutes
from src.config import MINUTES_DIR


# ---------- Markdown output ----------

def format_as_markdown(minutes: MeetingMinutes) -> str:
    date_str = datetime.now().strftime("%B %d, %Y")

    md = []
    md.append(f"# {minutes.title}")
    md.append(f"*Generated on {date_str}*\n")

    if minutes.attendees:
        md.append("## Attendees")
        for person in minutes.attendees:
            md.append(f"- {person}")
        md.append("")

    md.append("## Summary")
    md.append(minutes.summary)
    md.append("")

    md.append("## Key Decisions")
    if minutes.key_decisions:
        for decision in minutes.key_decisions:
            md.append(f"- {decision}")
    else:
        md.append("- No formal decisions recorded.")
    md.append("")

    md.append("## Action Items")
    if minutes.action_items:
        md.append("| Task | Owner | Deadline |")
        md.append("|------|-------|----------|")
        for item in minutes.action_items:
            md.append(f"| {item.task} | {item.owner} | {item.deadline} |")
    else:
        md.append("No action items identified.")

    return "\n".join(md)


# ---------- Word (.docx) output ----------

def format_as_docx(minutes: MeetingMinutes, output_path: str):
    doc = Document()

    title = doc.add_heading(minutes.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    date_str = datetime.now().strftime("%B %d, %Y")
    date_para = doc.add_paragraph(f"Generated on {date_str}")
    date_para.runs[0].italic = True

    if minutes.attendees:
        doc.add_heading("Attendees", level=1)
        for person in minutes.attendees:
            doc.add_paragraph(person, style="List Bullet")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(minutes.summary)

    doc.add_heading("Key Decisions", level=1)
    if minutes.key_decisions:
        for decision in minutes.key_decisions:
            doc.add_paragraph(decision, style="List Bullet")
    else:
        doc.add_paragraph("No formal decisions recorded.")

    doc.add_heading("Action Items", level=1)
    if minutes.action_items:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Task"
        hdr_cells[1].text = "Owner"
        hdr_cells[2].text = "Deadline"

        for item in minutes.action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.task
            row_cells[1].text = item.owner
            row_cells[2].text = item.deadline
    else:
        doc.add_paragraph("No action items identified.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Word doc saved to {output_path}")


# ---------- Public helper ----------

def save_minutes(minutes: MeetingMinutes, base_filename: str, formats=("md", "docx")) -> dict:
    """
    Saves minutes in requested formats.
    Returns dict of {format: filepath}
    """
    os.makedirs(MINUTES_DIR, exist_ok=True)
    saved = {}

    if "md" in formats:
        md_content = format_as_markdown(minutes)
        md_path = os.path.join(MINUTES_DIR, f"{base_filename}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        saved["md"] = md_path
        print(f"Markdown saved to {md_path}")

    if "docx" in formats:
        docx_path = os.path.join(MINUTES_DIR, f"{base_filename}.docx")
        format_as_docx(minutes, docx_path)
        saved["docx"] = docx_path

    return saved


if __name__ == "__main__":
    # quick manual test using summarizer's sample
    from src.summarizer import summarize_transcript

    sample_transcript = """
    John: Alright team, let's kick off. We need to finalize the Q3 roadmap.
    Priya: I think we should prioritize the mobile app redesign first.
    John: Agreed. Priya, can you have the wireframes ready by next Friday?
    Priya: Yes, I'll get that done.
    Raj: I'll handle backend API changes, should take about two weeks.
    John: Great, let's also decide — we're dropping the legacy dashboard feature.
    Priya: Confirmed, removing it from this quarter's scope.
    John: Perfect. Let's reconvene next week.
    """

    minutes = summarize_transcript(sample_transcript)
    save_minutes(minutes, base_filename="test_meeting")